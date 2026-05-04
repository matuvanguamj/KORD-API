"""
KORD DOCX — Rapport Word éditable
Style JLL premium — Contenu complet et bien formaté
"""
import io
from datetime import datetime
from typing import Dict, Any

from docx import Document
from docx.shared import Pt, Mm, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.replace('#',''))
    tcPr.append(shd)


def add_border_bottom(para, color='000000', size='6'):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), size)
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def heading(doc, num, title, subtitle=''):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(2)
    r_num = p.add_run(f"{num}  ")
    r_num.font.size = Pt(8)
    r_num.font.bold = True
    r_num.font.color.rgb = RGBColor(0x99,0x99,0x99)
    r_title = p.add_run(title.upper())
    r_title.font.size = Pt(11)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(0,0,0)
    add_border_bottom(p, '000000', '8')
    if subtitle:
        ps = doc.add_paragraph()
        rs = ps.add_run(subtitle)
        rs.font.size = Pt(9)
        rs.font.color.rgb = RGBColor(0x88,0x88,0x88)
        rs.font.italic = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def stat_row(doc, stats):
    """Ligne de stats style JLL — grand nombre + label."""
    t = doc.add_table(rows=1, cols=len(stats))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (num, label) in enumerate(stats):
        cell = t.rows[0].cells[i]
        set_cell_bg(cell, '000000')
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(str(num))
        r1.font.size = Pt(28)
        r1.font.bold = True
        r1.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(label.upper())
        r2.font.size = Pt(7)
        r2.font.color.rgb = RGBColor(0xAA,0xAA,0xAA)
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(8)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def pull_quote(doc, text):
    """Citation mise en valeur — style JLL."""
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    bar_cell = t.rows[0].cells[0]
    set_cell_bg(bar_cell, '000000')
    bar_cell.width = Mm(4)
    text_cell = t.rows[0].cells[1]
    set_cell_bg(text_cell, 'F4F3F0')
    p = text_cell.paragraphs[0]
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0,0,0)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Mm(4)
    p.paragraph_format.right_indent = Mm(4)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def body_text(doc, text, bold_prefix=None, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Mm(6)
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        rb = p.add_run(f"{bold_prefix} ")
        rb.font.bold = True
        rb.font.size = Pt(9.5)
        rb.font.color.rgb = RGBColor(0,0,0)
    r = p.add_run(text)
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x1A,0x1A,0x1A)
    return p


def add_image(doc, img_bytes, width_cm):
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(io.BytesIO(img_bytes), width=Cm(width_cm))
    except Exception as e:
        print(f"Image error: {e}")


def source_caption(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(f"Source : {text}")
    r.font.size = Pt(7)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0xAA,0xAA,0xAA)
    p.paragraph_format.space_after = Pt(6)


def generate_prereport_docx(
    consolidated, recommendations, all_results,
    client_name="Client", company_name="",
    trimestre="", gauge_png=None, bar_png=None,
    radar_png=None, evol_png=None,
) -> bytes:

    doc = Document()
    for section in doc.sections:
        section.page_width  = Mm(297)
        section.page_height = Mm(210)
        section.top_margin    = Mm(18)
        section.bottom_margin = Mm(18)
        section.left_margin   = Mm(18)
        section.right_margin  = Mm(18)

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    score    = consolidated.get("score_total", 0)
    now      = datetime.now()
    date_str = now.strftime("%d %B %Y")
    trim_str = trimestre or f"T{(now.month-1)//3+1} {now.year}"
    nb       = len(all_results)
    reco     = recommendations

    # ── COUVERTURE ──
    # Header noir
    t_cover = doc.add_table(rows=1, cols=2)
    t_cover.alignment = WD_TABLE_ALIGNMENT.LEFT

    left = t_cover.rows[0].cells[0]
    set_cell_bg(left, '000000')
    left.width = Mm(130)
    p_kord = left.paragraphs[0]
    r_kord = p_kord.add_run("KORD")
    r_kord.font.size = Pt(48)
    r_kord.font.bold = True
    r_kord.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    p_kord.paragraph_format.space_before = Pt(14)

    p_sub = left.add_paragraph()
    r_sub = p_sub.add_run("AUDIT DE PERFORMANCE OPÉRATIONNELLE")
    r_sub.font.size = Pt(8)
    r_sub.font.color.rgb = RGBColor(0x77,0x77,0x77)

    left.add_paragraph()

    p_co = left.add_paragraph()
    r_co = p_co.add_run((company_name or client_name).upper())
    r_co.font.size = Pt(20)
    r_co.font.bold = True
    r_co.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

    if company_name and client_name:
        p_cl = left.add_paragraph()
        r_cl = p_cl.add_run(client_name)
        r_cl.font.size = Pt(11)
        r_cl.font.color.rgb = RGBColor(0xBB,0xBB,0xBB)

    p_meta = left.add_paragraph()
    r_meta = p_meta.add_run(f"{trim_str}  ·  {date_str}  ·  {nb} fichier{'s' if nb>1 else ''} analysé{'s' if nb>1 else ''}")
    r_meta.font.size = Pt(8)
    r_meta.font.color.rgb = RGBColor(0x66,0x66,0x66)
    p_meta.paragraph_format.space_before = Pt(8)
    p_meta.paragraph_format.space_after = Pt(14)

    right = t_cover.rows[0].cells[1]
    set_cell_bg(right, 'F4F3F0')
    p_sc_lbl = right.paragraphs[0]
    p_sc_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sc_lbl = p_sc_lbl.add_run("SCORE KORD")
    r_sc_lbl.font.size = Pt(9)
    r_sc_lbl.font.bold = True
    r_sc_lbl.font.color.rgb = RGBColor(0x88,0x88,0x88)
    p_sc_lbl.paragraph_format.space_before = Pt(14)

    p_sc = right.add_paragraph()
    p_sc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sc = p_sc.add_run(str(score))
    r_sc.font.size = Pt(64)
    r_sc.font.bold = True
    r_sc.font.color.rgb = RGBColor(0,0,0)

    p_sc2 = right.add_paragraph()
    p_sc2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sc2 = p_sc2.add_run("/100  —  INDICATIF")
    r_sc2.font.size = Pt(8)
    r_sc2.font.color.rgb = RGBColor(0x88,0x88,0x88)

    if gauge_png:
        try:
            p_g = right.add_paragraph()
            p_g.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_g.add_run().add_picture(io.BytesIO(gauge_png), width=Cm(8))
        except: pass

    doc.add_paragraph()

    # Note travail
    t_note = doc.add_table(rows=1, cols=1)
    nc = t_note.rows[0].cells[0]
    set_cell_bg(nc, 'F4F3F0')
    p_note = nc.paragraphs[0]
    r_note = p_note.add_run("⚠  DOCUMENT DE TRAVAIL — USAGE INTERNE KORD UNIQUEMENT — Ne pas transmettre au client avant validation")
    r_note.font.size = Pt(8)
    r_note.font.bold = True
    r_note.font.color.rgb = RGBColor(0x66,0x66,0x66)
    p_note.paragraph_format.space_before = Pt(6)
    p_note.paragraph_format.space_after = Pt(6)
    doc.add_paragraph()

    # ── 01 RÉSUMÉ EXÉCUTIF ──
    heading(doc, "01", "Résumé Exécutif", "Situation globale et principaux constats")

    # Stats visuelles
    alertes_n = len(consolidated.get("alertes",[]))
    prios_n   = len(reco.get("priorites",[]))
    stat_row(doc, [
        (f"{score}/100", "Score KORD indicatif"),
        (str(alertes_n), f"Anomalie{'s' if alertes_n>1 else ''} détectée{'s' if alertes_n>1 else ''}"),
        (str(prios_n), f"Priorité{'s' if prios_n>1 else ''} d'action"),
        (str(nb), f"Fichier{'s' if nb>1 else ''} analysé{'s' if nb>1 else ''}"),
    ])

    # Résumé paragraphes
    resume = reco.get("resume_executif", {})
    if isinstance(resume, dict):
        for key in ["paragraphe_1","paragraphe_2","paragraphe_3"]:
            txt = resume.get(key,"")
            if txt and txt not in ["...", ""]:
                body_text(doc, txt)
    elif isinstance(resume, str) and resume:
        body_text(doc, resume)

    doc.add_paragraph()

    # Message dirigeant — pull quote
    msg = reco.get("message_dirigeant","")
    if msg and msg not in ["...", ""]:
        pull_quote(doc, msg)

    # Benchmark
    bench = reco.get("benchmark","")
    if bench and bench not in ["...", ""]:
        t_bench = doc.add_table(rows=1, cols=2)
        bc1 = t_bench.rows[0].cells[0]
        bc2 = t_bench.rows[0].cells[1]
        set_cell_bg(bc1, 'F4F3F0')
        set_cell_bg(bc2, 'F4F3F0')
        bc1.width = Mm(35)
        p_bl = bc1.paragraphs[0]
        r_bl = p_bl.add_run("BENCHMARK\nSECTORIEL")
        r_bl.font.size = Pt(7)
        r_bl.font.bold = True
        r_bl.font.color.rgb = RGBColor(0x88,0x88,0x88)
        p_bl.paragraph_format.space_before = Pt(6)
        p_bl.paragraph_format.space_after = Pt(6)
        p_bv = bc2.paragraphs[0]
        r_bv = p_bv.add_run(bench)
        r_bv.font.size = Pt(9)
        r_bv.font.color.rgb = RGBColor(0x1A,0x1A,0x1A)
        p_bv.paragraph_format.space_before = Pt(6)
        p_bv.paragraph_format.space_after = Pt(6)
    doc.add_paragraph()

    # ── 02 GRAPHIQUES ──
    heading(doc, "02", "Analyse Graphique", "Score KORD vs Benchmark marché")

    if bar_png and radar_png:
        t_charts = doc.add_table(rows=1, cols=2)
        t_charts.alignment = WD_TABLE_ALIGNMENT.CENTER
        for img, ci in [(bar_png, 0), (radar_png, 1)]:
            try:
                cell = t_charts.rows[0].cells[ci]
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(io.BytesIO(img), width=Cm(11 if ci==0 else 8))
            except: pass
    elif bar_png:
        add_image(doc, bar_png, 22)
    elif radar_png:
        add_image(doc, radar_png, 14)

    if bar_png or radar_png:
        source_caption(doc, "Moteur d'analyse KORD — Données client — Benchmark PME secteur")

    if evol_png:
        add_image(doc, evol_png, 24)
        source_caption(doc, "Trajectoire indicative — sera enrichie à chaque trimestre")

    doc.add_paragraph()

    # ── 03 ANALYSE PAR PILIER ──
    heading(doc, "03", "Performance par Pilier", "Analyse détaillée avec chiffres et risques")

    pilier_order = ["stock_cash","transport_service","achats_fournisseurs","marges_retours","donnees_pilotage"]
    analyse_piliers = reco.get("analyse_piliers", {})
    pilier_labels = {
        "stock_cash":          ("Stock et Cash Immobilisé",     30),
        "transport_service":   ("Transport et Taux de Service", 20),
        "achats_fournisseurs": ("Achats et Fournisseurs",       20),
        "marges_retours":      ("Marges et Retours Clients",    15),
        "donnees_pilotage":    ("Données et Pilotage",          15),
    }

    for key in pilier_order:
        dp = analyse_piliers.get(key, {})
        label, max_pts = pilier_labels.get(key, (key, 20))
        sc = consolidated.get("analyses",{}).get(key,{}).get("score",0)
        if isinstance(dp, dict):
            sc = dp.get("score", sc)
            titre = dp.get("titre", label)
            niveau = dp.get("niveau","moyen").upper()
            analyse = dp.get("analyse","")
            chiffres = dp.get("chiffres_cles",[])
            risque = dp.get("risque_principal","")
        else:
            titre = label
            niveau = "MOYEN"
            analyse = chiffres = risque = ""
        pct = round(sc/max_pts*100) if max_pts > 0 else 0
        bg = '000000' if pct < 45 else '1A1A1A' if pct < 70 else '444444'

        # Header pilier
        t_ph = doc.add_table(rows=1, cols=3)
        c0 = t_ph.rows[0].cells[0]
        c1 = t_ph.rows[0].cells[1]
        c2 = t_ph.rows[0].cells[2]
        c0.width = Mm(120)
        c1.width = Mm(25)
        c2.width = Mm(20)
        for cell in [c0,c1,c2]:
            set_cell_bg(cell, bg)
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(titre.upper())
        r0.font.size = Pt(9)
        r0.font.bold = True
        r0.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        p0.paragraph_format.space_before = Pt(6)
        p0.paragraph_format.space_after = Pt(6)
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(f"{sc}/{max_pts}")
        r1.font.size = Pt(13)
        r1.font.bold = True
        r1.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_before = Pt(4)
        p2 = c2.paragraphs[0]
        r2 = p2.add_run(niveau)
        r2.font.size = Pt(7)
        r2.font.bold = True
        r2.font.color.rgb = RGBColor(0xCC,0xCC,0xCC)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(8)

        # Analyse narrative
        if analyse and analyse not in ["...",""]:
            p_an = doc.add_paragraph()
            p_an.paragraph_format.left_indent = Mm(0)
            p_an.paragraph_format.space_after = Pt(2)
            r_an = p_an.add_run(analyse)
            r_an.font.size = Pt(9.5)
            r_an.font.color.rgb = RGBColor(0x1A,0x1A,0x1A)

        # Chiffres clés
        for cf in chiffres:
            if cf and cf not in ["...", ""]:
                p_cf = doc.add_paragraph()
                p_cf.paragraph_format.left_indent = Mm(4)
                p_cf.paragraph_format.space_after = Pt(1)
                r_cf = p_cf.add_run(f"▸  {cf}")
                r_cf.font.size = Pt(9)
                r_cf.font.color.rgb = RGBColor(0x44,0x44,0x44)

        # Risque principal
        if risque and risque not in ["...", ""]:
            t_rq = doc.add_table(rows=1, cols=2)
            rq0 = t_rq.rows[0].cells[0]
            rq1 = t_rq.rows[0].cells[1]
            rq0.width = Mm(30)
            set_cell_bg(rq0, 'F4F3F0')
            set_cell_bg(rq1, 'F4F3F0')
            p_rql = rq0.paragraphs[0]
            r_rql = p_rql.add_run("RISQUE")
            r_rql.font.size = Pt(7)
            r_rql.font.bold = True
            r_rql.font.color.rgb = RGBColor(0x88,0x88,0x88)
            p_rql.paragraph_format.space_before = Pt(4)
            p_rql.paragraph_format.space_after = Pt(4)
            p_rqv = rq1.paragraphs[0]
            r_rqv = p_rqv.add_run(risque)
            r_rqv.font.size = Pt(9)
            r_rqv.font.bold = True
            r_rqv.font.color.rgb = RGBColor(0,0,0)
            p_rqv.paragraph_format.space_before = Pt(4)
            p_rqv.paragraph_format.space_after = Pt(4)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── 04 CROISEMENTS CLÉS ──
    croisements = reco.get("croisements_cles", [])
    if croisements:
        heading(doc, "04", "Croisements Clés", "Ce que révèle l'analyse croisée de vos données")
        for cr in croisements:
            if not isinstance(cr, dict): continue
            fichiers = cr.get("fichiers",[])
            obs = cr.get("observation","")
            impact = cr.get("impact","")
            if obs and obs not in ["...",""]:
                t_cr = doc.add_table(rows=1, cols=3)
                set_cell_bg(t_cr.rows[0].cells[0], 'F4F3F0')
                set_cell_bg(t_cr.rows[0].cells[1], 'FFFFFF')
                set_cell_bg(t_cr.rows[0].cells[2], 'F4F3F0')
                t_cr.rows[0].cells[0].width = Mm(42)
                t_cr.rows[0].cells[2].width = Mm(50)
                p_f = t_cr.rows[0].cells[0].paragraphs[0]
                r_f = p_f.add_run(" × ".join(fichiers))
                r_f.font.size = Pt(7.5)
                r_f.font.bold = True
                r_f.font.color.rgb = RGBColor(0x44,0x44,0x44)
                p_f.paragraph_format.space_before = Pt(5)
                p_f.paragraph_format.space_after = Pt(5)
                p_o = t_cr.rows[0].cells[1].paragraphs[0]
                r_o = p_o.add_run(obs)
                r_o.font.size = Pt(9)
                r_o.font.color.rgb = RGBColor(0x1A,0x1A,0x1A)
                p_o.paragraph_format.space_before = Pt(5)
                p_o.paragraph_format.space_after = Pt(5)
                p_i = t_cr.rows[0].cells[2].paragraphs[0]
                r_i = p_i.add_run(impact)
                r_i.font.size = Pt(9)
                r_i.font.bold = True
                r_i.font.color.rgb = RGBColor(0,0,0)
                p_i.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p_i.paragraph_format.space_before = Pt(5)
                p_i.paragraph_format.space_after = Pt(5)
        doc.add_paragraph()

    # ── 05 ANOMALIES ──
    anomalies = reco.get("anomalies", consolidated.get("alertes",[]))
    if anomalies:
        heading(doc, "05", "Anomalies Détectées", "Classées par urgence — Impact business estimé")

        t_an = doc.add_table(rows=1, cols=4)
        hdrs = ["URGENCE","ANOMALIE DÉTECTÉE","IMPACT BUSINESS","IMPACT €"]
        for i, h in enumerate(hdrs):
            cell = t_an.rows[0].cells[i]
            set_cell_bg(cell, '000000')
            p = cell.paragraphs[0]
            r = p.add_run(h)
            r.font.size = Pt(7)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
        t_an.rows[0].cells[0].width = Mm(20)
        t_an.rows[0].cells[2].width = Mm(60)
        t_an.rows[0].cells[3].width = Mm(35)

        u_order = {"CRITIQUE":0,"MOYEN":1,"FAIBLE":2}
        if anomalies and isinstance(anomalies[0], dict) and "urgence" in anomalies[0]:
            anomalies = sorted(anomalies, key=lambda x: u_order.get(x.get("urgence","FAIBLE"),2))

        for i, al in enumerate(anomalies[:12]):
            if not isinstance(al, dict): continue
            row = t_an.add_row()
            bg = 'FFFFFF' if i % 2 == 0 else 'F9F9F9'
            for cell in row.cells:
                set_cell_bg(cell, bg)
            row.cells[0].width = Mm(20)
            row.cells[2].width = Mm(60)
            row.cells[3].width = Mm(35)

            urg = al.get("urgence","MOYEN")
            urg_bg = '000000' if urg=="CRITIQUE" else '444444' if urg=="MOYEN" else '888888'
            set_cell_bg(row.cells[0], urg_bg)
            p0 = row.cells[0].paragraphs[0]
            r0 = p0.add_run(urg)
            r0.font.size = Pt(6.5)
            r0.font.bold = True
            r0.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p0.paragraph_format.space_before = Pt(3)
            p0.paragraph_format.space_after = Pt(3)

            titre = al.get("titre", al.get("message",""))
            detect = al.get("detection", al.get("message",""))
            p1 = row.cells[1].paragraphs[0]
            r1t = p1.add_run(titre + "\n")
            r1t.font.size = Pt(9)
            r1t.font.bold = True
            r1t.font.color.rgb = RGBColor(0,0,0)
            r1d = p1.add_run(detect)
            r1d.font.size = Pt(8.5)
            r1d.font.color.rgb = RGBColor(0x33,0x33,0x33)
            p1.paragraph_format.space_before = Pt(3)
            p1.paragraph_format.space_after = Pt(3)

            p2 = row.cells[2].paragraphs[0]
            r2 = p2.add_run(al.get("impact_business",""))
            r2.font.size = Pt(8.5)
            r2.font.color.rgb = RGBColor(0x1A,0x1A,0x1A)
            p2.paragraph_format.space_before = Pt(3)
            p2.paragraph_format.space_after = Pt(3)

            p3 = row.cells[3].paragraphs[0]
            r3 = p3.add_run(al.get("impact_financier",""))
            r3.font.size = Pt(9)
            r3.font.bold = True
            r3.font.color.rgb = RGBColor(0,0,0)
            p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p3.paragraph_format.space_before = Pt(3)
            p3.paragraph_format.space_after = Pt(3)

        doc.add_paragraph()

    # ── 06 PRIORITÉS D'ACTION ──
    priorites = reco.get("priorites",[])
    if priorites:
        heading(doc, "06", "Priorités d'Action", "Top actions classées par impact financier")

        for p_item in priorites[:5]:
            if not isinstance(p_item, dict): continue
            rang       = p_item.get("rang","")
            titre      = p_item.get("titre","")
            probleme   = p_item.get("probleme","")
            action     = p_item.get("action","")
            impact_att = p_item.get("impact_attendu","")
            gain       = p_item.get("gain_potentiel","")
            delai      = p_item.get("delai","")
            complexite = p_item.get("complexite","")
            qw         = p_item.get("quick_win", False)

            # Header priorité
            t_phead = doc.add_table(rows=1, cols=3)
            ph0 = t_phead.rows[0].cells[0]
            ph1 = t_phead.rows[0].cells[1]
            ph2 = t_phead.rows[0].cells[2]
            ph0.width = Mm(14)
            ph2.width = Mm(28)
            set_cell_bg(ph0, 'F4F3F0')
            set_cell_bg(ph1, '000000')
            set_cell_bg(ph2, '1A1A1A' if qw else '000000')
            pp0 = ph0.paragraphs[0]
            rp0 = pp0.add_run(f"0{rang}")
            rp0.font.size = Pt(18)
            rp0.font.bold = True
            rp0.font.color.rgb = RGBColor(0xCC,0xCC,0xCC)
            pp0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pp0.paragraph_format.space_before = Pt(6)
            pp0.paragraph_format.space_after = Pt(6)
            pp1 = ph1.paragraphs[0]
            rp1 = pp1.add_run(titre.upper())
            rp1.font.size = Pt(10)
            rp1.font.bold = True
            rp1.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            pp1.paragraph_format.space_before = Pt(8)
            pp1.paragraph_format.space_after = Pt(8)
            pp2 = ph2.paragraphs[0]
            rp2 = pp2.add_run("⚡ QUICK WIN" if qw else delai)
            rp2.font.size = Pt(7.5)
            rp2.font.bold = True
            rp2.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            pp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pp2.paragraph_format.space_before = Pt(8)

            # Corps 2 colonnes
            t_pbody = doc.add_table(rows=1, cols=2)
            pb_left  = t_pbody.rows[0].cells[0]
            pb_right = t_pbody.rows[0].cells[1]
            set_cell_bg(pb_left, 'FAFAFA')
            set_cell_bg(pb_right, 'F4F3F0')
            pb_left.width = Mm(140)

            for lbl, val in [("Problème détecté", probleme), ("Action recommandée", action)]:
                if val and val not in ["...", ""]:
                    p_l = pb_left.add_paragraph()
                    r_lbl = p_l.add_run(f"{lbl} : ")
                    r_lbl.font.size = Pt(8.5)
                    r_lbl.font.bold = True
                    r_lbl.font.color.rgb = RGBColor(0,0,0)
                    r_val = p_l.add_run(val)
                    r_val.font.size = Pt(9)
                    r_val.font.color.rgb = RGBColor(0x1A,0x1A,0x1A)
                    p_l.paragraph_format.left_indent = Mm(4)
                    p_l.paragraph_format.space_before = Pt(3)
                    p_l.paragraph_format.space_after = Pt(3)

            for lbl, val, bold in [
                ("Impact attendu", impact_att, True),
                ("Gain estimé", gain, True),
                ("Complexité", complexite, False),
                ("Délai", delai, False),
            ]:
                if val and val not in ["...", ""]:
                    p_r = pb_right.add_paragraph()
                    r_rl = p_r.add_run(f"{lbl}\n")
                    r_rl.font.size = Pt(7.5)
                    r_rl.font.bold = True
                    r_rl.font.color.rgb = RGBColor(0x88,0x88,0x88)
                    r_rv = p_r.add_run(val)
                    r_rv.font.size = Pt(9.5 if bold else 8.5)
                    r_rv.font.bold = bold
                    r_rv.font.color.rgb = RGBColor(0,0,0)
                    p_r.paragraph_format.left_indent = Mm(4)
                    p_r.paragraph_format.space_before = Pt(4)
                    p_r.paragraph_format.space_after = Pt(4)

            doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── 07 OPPORTUNITÉS CACHÉES ──
    opps = reco.get("opportunites_cachees",[])
    if opps:
        heading(doc, "07", "Opportunités Cachées", "Leviers additionnels dans vos données")
        for opp in opps:
            if not isinstance(opp, dict): continue
            t_opp = doc.add_table(rows=1, cols=2)
            oc1 = t_opp.rows[0].cells[0]
            oc2 = t_opp.rows[0].cells[1]
            oc2.width = Mm(45)
            set_cell_bg(oc1, '1A1A1A')
            set_cell_bg(oc2, '1A1A1A')
            p_ot = oc1.paragraphs[0]
            r_ot = p_ot.add_run(opp.get("titre",""))
            r_ot.font.size = Pt(10)
            r_ot.font.bold = True
            r_ot.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            p_ot.paragraph_format.space_before = Pt(7)
            p_ot.paragraph_format.space_after = Pt(3)
            p_od = oc1.add_paragraph()
            r_od = p_od.add_run(opp.get("description",""))
            r_od.font.size = Pt(9)
            r_od.font.color.rgb = RGBColor(0xCC,0xCC,0xCC)
            p_od.paragraph_format.space_after = Pt(7)
            p_og = oc2.paragraphs[0]
            r_og = p_og.add_run(opp.get("gain_estime",""))
            r_og.font.size = Pt(13)
            r_og.font.bold = True
            r_og.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            p_og.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_og.paragraph_format.space_before = Pt(12)
        doc.add_paragraph()

    # ── 08 POINTS DE VIGILANCE ──
    vigilance = reco.get("points_vigilance",[])
    if vigilance:
        heading(doc, "08", "Points de Vigilance")
        for i, pt in enumerate(vigilance, 1):
            if pt and pt not in ["...",""]:
                p_v = doc.add_paragraph()
                r_vn = p_v.add_run(f"{i:02d}.  ")
                r_vn.font.bold = True
                r_vn.font.size = Pt(9)
                r_vn.font.color.rgb = RGBColor(0xBB,0xBB,0xBB)
                r_vt = p_v.add_run(pt)
                r_vt.font.size = Pt(9.5)
                r_vt.font.color.rgb = RGBColor(0x1A,0x1A,0x1A)
                p_v.paragraph_format.space_after = Pt(4)
        doc.add_paragraph()

    # ── 09 QUESTIONS RESTITUTION ──
    questions = reco.get("questions_restitution",[])
    if questions:
        heading(doc, "09", "Questions — Session de Restitution Client")
        t_q = doc.add_table(rows=len(questions), cols=2)
        for i, q in enumerate(questions):
            if not q or q in ["...",""]: continue
            qc0 = t_q.rows[i].cells[0]
            qc1 = t_q.rows[i].cells[1]
            qc0.width = Mm(12)
            set_cell_bg(qc0, '000000')
            set_cell_bg(qc1, 'F4F3F0' if i%2==0 else 'FFFFFF')
            pq0 = qc0.paragraphs[0]
            rq0 = pq0.add_run(f"Q{i+1}")
            rq0.font.size = Pt(8)
            rq0.font.bold = True
            rq0.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            pq0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pq0.paragraph_format.space_before = Pt(5)
            pq0.paragraph_format.space_after = Pt(5)
            pq1 = qc1.paragraphs[0]
            rq1 = pq1.add_run(q)
            rq1.font.size = Pt(9.5)
            rq1.font.color.rgb = RGBColor(0x1A,0x1A,0x1A)
            pq1.paragraph_format.space_before = Pt(5)
            pq1.paragraph_format.space_after = Pt(5)
        doc.add_paragraph()

    # ── 10 PROCHAINE ÉTAPE ──
    next_step = reco.get("prochaine_etape","")
    if next_step and next_step not in ["...",""]:
        heading(doc, "10", "Prochaine Étape — Session de Restitution KORD")
        pull_quote(doc, next_step)

    # ── ANNEXE LIBRE ──
    doc.add_paragraph()
    p_sep = doc.add_paragraph()
    add_border_bottom(p_sep, 'CCCCCC', '4')
    heading(doc, "—", "Annexe — Notes et Commentaires Experts KORD")
    p_ann = doc.add_paragraph()
    r_ann = p_ann.add_run("[ Espace réservé à l'équipe KORD pour ajouter observations, données complémentaires et recommandations avant restitution. ]")
    r_ann.font.size = Pt(9)
    r_ann.font.italic = True
    r_ann.font.color.rgb = RGBColor(0xBB,0xBB,0xBB)
    doc.add_paragraph()
    doc.add_paragraph()

    # ── FOOTER ──
    p_ft = doc.add_paragraph()
    add_border_bottom(p_ft, 'CCCCCC', '4')
    p_ft2 = doc.add_paragraph()
    r_ft = p_ft2.add_run(f"KORD — Audit de performance opérationnelle — {trim_str} — Document de travail confidentiel — Score indicatif avant validation experts KORD")
    r_ft.font.size = Pt(7)
    r_ft.font.color.rgb = RGBColor(0xAA,0xAA,0xAA)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
